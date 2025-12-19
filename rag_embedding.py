import os
import shutil
import pandas as pd
from docx import Document
from dotenv import load_dotenv

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_cohere import CohereEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document as LCDocument

load_dotenv()

# =====================
# CONFIG
# =====================
BASE_PATH = r"E:\Cursos\Gettalent_pi\LimonIA\datos"
PROCESADOS_PATH = r"E:\Cursos\Gettalent_pi\LimonIA\procesados"
CHROMA_PATH = r"E:\Cursos\Gettalent_pi\LimonIA\chroma_db"

COHERE_MODEL = "embed-multilingual-v3.0"
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150

# =====================
# CARGA DE ARCHIVOS
# =====================
def cargar_excel(path, filename):
    documentos = []

    hojas = pd.read_excel(path, sheet_name=None)

    for nombre_hoja, df in hojas.items():
        df = df.fillna("")

        print(f"   📄 Hoja: {nombre_hoja} ({len(df)} filas)")

        for _, row in df.iterrows():
            texto = " | ".join([f"{col}: {row[col]}" for col in df.columns])

            documentos.append(
                LCDocument(
                    page_content=texto,
                    metadata={
                        "source": filename,
                        "sheet": nombre_hoja
                    }
                )
            )

    return documentos


def cargar_docx(path, filename):
    documentos = []
    doc = Document(path)

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            documentos.append(
                LCDocument(
                    page_content=p.text.strip(),
                    metadata={
                        "source": filename,
                        "paragraph": i
                    }
                )
            )

    return documentos


def cargar_archivos(base_path):
    documentos = []
    archivos_procesados = []

    for archivo in os.listdir(base_path):
        ruta = os.path.join(base_path, archivo)

        if not os.path.isfile(ruta):
            continue

        if archivo.endswith(".xlsx"):
            print(f"📊 Leyendo Excel: {archivo}")
            documentos.extend(cargar_excel(ruta, archivo))
            archivos_procesados.append(ruta)

        elif archivo.endswith(".docx"):
            print(f"📄 Leyendo Word: {archivo}")
            documentos.extend(cargar_docx(ruta, archivo))
            archivos_procesados.append(ruta)

    return documentos, archivos_procesados


# =====================
# MAIN
# =====================
def main():
    documentos, archivos = cargar_archivos(BASE_PATH)

    if not documentos:
        print("⚠️ No hay archivos nuevos para procesar")
        return

    print(f"\n📚 Documentos cargados: {len(documentos)}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    docs = splitter.split_documents(documentos)

    print(f"✂️ Chunks creados: {len(docs)}")

    embeddings = CohereEmbeddings(
        model=COHERE_MODEL,
        cohere_api_key=os.getenv("COHERE_API_KEY")
    )

    # 🔹 ABRIMOS LA BASE EXISTENTE
    db = Chroma(
        persist_directory=CHROMA_PATH,
        embedding_function=embeddings
    )

    # 🔹 AGREGAMOS (NO PISAMOS)
    db.add_documents(docs)
    db.persist()

    print("✅ Embeddings agregados correctamente")

    # 🔹 MOVER LOS ARCHIVOS YA PROCESADOS
    for archivo in archivos:
        try:
            # Crear la carpeta 'procesados' si no existe
            if not os.path.exists(PROCESADOS_PATH):
                os.makedirs(PROCESADOS_PATH)

            # Mover archivo
            archivo_destino = os.path.join(PROCESADOS_PATH, os.path.basename(archivo))
            shutil.move(archivo, archivo_destino)

            print(f"✅ Archivo movido: {os.path.basename(archivo)} a {PROCESADOS_PATH}")

        except Exception as e:
            print(f"⚠️ No se pudo mover {archivo}: {e}")


if __name__ == "__main__":
    main()
